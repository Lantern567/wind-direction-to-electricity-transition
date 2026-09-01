"""
把 v3.4 正文图2 槽位换成 wp9f 重绘图（27/171 新口径）
- 位置法：图2 图注段落【之前】第一个含 w:drawing 的段落（本稿图片在图注之前）
- 宽度保持原图2 的 Cm 设置；Word 占用 → _tmp 回退
"""
import os, io, sys
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Cm
from docx.oxml.ns import qn

REPO = r'D:\01学习资料\wind-direction-to-electricity-transition'
SRC = os.path.join(REPO, '四场景风速风向分解贡献', 'Nature_Energy_中文稿_论点与论据框架_v3.4.docx')
PNG = os.path.join(REPO, '结论三、四重算结果', 'figures', 'fig2_v34.png')
assert os.path.exists(PNG)

doc = Document(SRC)
paras = doc.paragraphs

cap_i = next(i for i, p in enumerate(paras) if p.text.startswith('图2｜'))
img_i = None
for j in range(cap_i - 1, -1, -1):
    if paras[j]._element.findall('.//' + qn('w:drawing')):
        img_i = j
        break
assert img_i is not None, '图2 图注前未找到图片段落'
print(f'图2 图注段落 {cap_i}，图片段落 {img_i}')

# 原图宽度（extent cx / 914400 = 英寸）
p = paras[img_i]
old_cm = None
for ext in p._element.findall('.//' + qn('wp:extent')):
    old_cm = int(ext.get('cx')) / 914400.0 * 2.54
print(f'原图2 宽度: {old_cm:.2f} cm')

for d in p._element.findall('.//' + qn('w:drawing')):
    d.getparent().remove(d)
run = p.add_run()
run.add_picture(PNG, width=Cm(old_cm if old_cm else 15.0))
print(f'新图2 ({os.path.getsize(PNG)//1024} KB) 已换入，宽度 {old_cm if old_cm else 15.0:.2f} cm')

try:
    doc.save(SRC)
    print('已保存 →', SRC)
except PermissionError:
    tmp = SRC.replace('.docx', '_新图2_tmp.docx')
    doc.save(tmp)
    print('原文件被 Word 占用，已存到:', tmp, '（关闭 Word 后改名回正式文件名）')
