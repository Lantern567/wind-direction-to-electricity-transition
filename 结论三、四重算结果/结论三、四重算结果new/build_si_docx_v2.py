# -*- coding: utf-8 -*-
"""补充信息（Supplementary Information）生成脚本 v2（2026-08-27）
====================================================================
v1 基础上按学长反馈扩充：
  1. 图件全部英文 + 字体统一（主稿图 1–4 的 NC 体例 / Arial 字号 8）：
     图 S1 重建为英文版（figures-nat/FigS1_paradigm_layouts.png）；
  2. 新增图 S2 基础数据总览（figures-nat/FigS2_basic_data.png）：
     171 场三区分布地图 + 容量/机组/水深/面积分布 + Weibull k–平均风速；
  3. 新增图 S3 敏感性分析可视化（figures-nat/FigS3_sensitivity.png）：
     记录长度敏感性 ECDF、机型口径敏感性（表 S8）、受控气候反事实
     C0–C3（表 S4）、6 场风速窗口机制分解（5.8 节补算 wp9g，口径 B）；
  4. 同步主稿 v3.11 口径：表 S1 越线项目 5 场→6 场（含 F160 美国东海岸）、
     5/108=4.6%、6/146、比值 1.03–2.85。
数值来源：冻结管线报告（wp5c/wp5d/wp6c/wp7c/wp7d/wp9c/wp9d/wp9g），
与主稿 v3.11 一致。
输出：结论三、四重算结果/补充信息_Supplementary_Information_v2.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUT = r'd:\01学习资料\wind-direction-to-electricity-transition\结论三、四重算结果\补充信息_Supplementary_Information_v2.docx'
FIG_S1 = r'd:\01学习资料\wind-direction-to-electricity-transition\结论三、四重算结果\figures-nat\FigS1_paradigm_layouts.png'
FIG_S2 = r'd:\01学习资料\wind-direction-to-electricity-transition\结论三、四重算结果\figures-nat\FigS2_basic_data.png'
FIG_S3 = r'd:\01学习资料\wind-direction-to-electricity-transition\结论三、四重算结果\figures-nat\FigS3_sensitivity.png'

doc = Document()

# ── 页面：A4 ──
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
sec.top_margin = sec.bottom_margin = Cm(2.54)
sec.left_margin = sec.right_margin = Cm(2.54)


def set_font(r, eastasia='宋体', ascii_font='Times New Roman', size=10.5,
             bold=None, color=None, italic=None):
    r.font.name = ascii_font
    r._element.rPr.rFonts.set(qn('w:eastAsia'), eastasia)
    r.font.size = Pt(size)
    if bold is not None:
        r.font.bold = bold
    if italic is not None:
        r.font.italic = italic
    if color is not None:
        r.font.color.rgb = RGBColor(*color)


def para(text, size=10.5, bold=False, align=None, eastasia='宋体',
         space_after=6, indent=None):
    p = doc.add_paragraph()
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if indent:
        p.paragraph_format.first_line_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    set_font(r, eastasia=eastasia, size=size, bold=bold)
    return p


def eq_para(segments, size=10.5, space_after=6):
    """segments: list of (text, mode), mode in {None, 'sub', 'sup', 'i'}"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    for text, mode in segments:
        r = p.add_run(text)
        set_font(r, size=size)
        if mode == 'sub':
            r.font.subscript = True
        elif mode == 'sup':
            r.font.superscript = True
        elif mode == 'i':
            r.font.italic = True
    return p


def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_font(r, eastasia='黑体', size=14, bold=True, color=(0, 0, 0))
    return p


def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, eastasia='黑体', size=12, bold=True, color=(0, 0, 0))
    return p


def table_caption(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, size=10.5, bold=True)
    return p


def add_table(headers, rows, col_widths=None, font_size=9, note=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = ''
        r = cell.paragraphs[0].add_run(h)
        set_font(r, size=font_size, bold=True)
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            cell = t.rows[i + 1].cells[j]
            cell.text = ''
            r = cell.paragraphs[0].add_run(str(v))
            set_font(r, size=font_size)
    if col_widths:
        for j, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[j].width = Cm(w)
    if note:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(note)
        set_font(r, size=9)
        r.font.italic = True
    else:
        doc.paragraphs[-1].paragraph_format.space_after = Pt(8)
    return t


def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    set_font(r, size=9)
    return p


# ══════════════ 封面 ══════════════
para('Supplementary Information', size=26, bold=True, align='center',
     space_after=4)
para('补充信息', size=16, bold=True, align='center', eastasia='黑体', space_after=24)
para('《密集排布与轴向风气候共同放大海上风电的朝向敏感性》', size=15, bold=True,
     align='center', eastasia='黑体', space_after=6)
para('Dense layouts and axial wind climates jointly amplify the orientation '
     'sensitivity of offshore wind farms', size=12, align='center', space_after=24)
para('作者：〔作者姓名，与主稿一致〕', size=12, align='center', space_after=4)
para('单位：〔单位名称，与主稿一致〕', size=12, align='center', space_after=4)
para('通讯作者：〔姓名，邮箱，与主稿一致〕', size=12, align='center', space_after=24)
para('本文件为上述主稿的补充信息，提供主稿未尽的样本与方法细节、指标定义，以及支撑正文 '
     '2.1–2.4 全部结论的补充表格与图件。文件内所有数值均由冻结计算管线生成'
     '（2014–2024 年逐时 ERA5 气象 × FLORIS v4.6 统一尾流回测），与主稿数字严格一致；'
     '逐场派生数据与计算脚本将在论文接收时随数据可用性声明一并公开。', size=10.5,
     align='center', space_after=0)
doc.add_page_break()

# ══════════════ 目录 ══════════════
h1('目录（Contents）')
toc = [
    'S1  补充方法',
    'S2  朝向敏感性长尾与越线风场（支撑正文 2.1）',
    'S3  高朝向响应的机制证据（支撑正文 2.2）',
    'S4  建设范式与全圆最优（支撑正文 2.3）',
    'S5  方向性走廊与项目级平行反事实（支撑正文 2.4）',
    'S6  敏感性分析可视化',
    'S7  统计口径与诚实边界',
    '补充参考文献（Supplementary References）',
    '',
    '表 S1–S12  补充表格',
    '图 S1  六套标准化建设范式的排布示意图（英文版）',
    '图 S2  基础数据总览（三区分布、样本组成、风况统计）',
    '图 S3  敏感性分析可视化（记录长度、机型、气候反事实、风速窗口）',
    '公式 (S1)–(S4)  指标定义',
]
for line in toc:
    para(line, size=11, space_after=3)
doc.add_page_break()

# ══════════════ S1 补充方法 ══════════════
h1('S1  补充方法')

h2('S1.1  研究样本与数据来源')
para('研究以全球真实建成海上风场的公开机位坐标为几何输入，覆盖 15,106 台机组'
     '（主稿参考文献[14]）。逐时风速与风向取 100 m 高度 ERA5 再分析数据'
     '（主稿参考文献[15]），并考虑再分析风电模拟的区域偏差边界（主稿参考文献[16]）。'
     '主评价时段为 2014–2024 年，1981–2010 年历史气候用于事前选角与反事实分解。'
     '所有风场采用统一 10 MW 参考机型（IEA Wind Task 37，D = 198 m）及一致的功率、'
     '推力曲线与电气损耗设置，以隔离排布与风况差异。按逐年累计投运机组构建'
     '风场—年样本后，得到 171 个风场和 1,203 个风场—年。', indent=0.74)
para('样本基础数据总览见图 S2：171 个风场分布于东亚（88 场）、欧洲（78 场）与美国'
     '东海岸（5 场）；全球样本以中国近海项目为主体，装机容量、机组数、场区面积呈'
     '长尾分布，水深中位数 14.5 m。场级多年平均风速（2014–2024，ERA5 100 m）为 '
     '5.75–11.20 m s⁻¹，Weibull 形状参数 k 为 2.01–3.28——低风速近海场址风况'
     '更稳定（k 更高）。', indent=0.74)
doc.add_picture(FIG_S2, width=Cm(15.0))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
caption('图 S2｜基础数据总览（图内文字为英文，字体与主稿图 1–4 一致）。(a) 171 场三区'
        '分布（红星为 6 个越线场 G > SD(MS)）；(b–e) 装机容量、机组数、水深、场区面积'
        '分布（柱上注记 n 与中位数）；(f) 场级 Weibull 形状参数 k 与 2014–2024 年'
        '平均风速，按区域着色。k 由场级风速边际直方图矩法拟合。')

h2('S1.2  统一模拟框架')
para('尾流计算采用 FLORIS v4.6 中的高斯解析尾流框架（主稿参考文献[17]）；'
     '方向—风速分箱预计算后，在逐时气象序列中查表积分。旋转测试、机位打乱对照'
     '和数值精度检查用于核对刚性旋转与能量积分的一致性。本文的“旋转”始终指整片'
     '阵列绕质心的刚性旋转，不改变机组、两两距离和内部几何；它既不是单机偏航，'
     '也不是运行期尾流偏转控制。', indent=0.74)

h2('S1.3  指标定义与口径（公式 S1–S4）')
para('以下四个指标与主稿 5.2–5.5 节的定义逐字一致，是全部分析的统一定义基础。',
     indent=0.74)
para('公式 (S1)：朝向增益 G——建成朝向相对历史气候最优朝向的模型化发电增量',
     size=10.5, bold=True, space_after=2)
eq_para([('G', 'i'), ('i,y', 'sub'), (' = [ E', None), ('i,y', 'sub'), ('(θ', None),
         ('i', 'sub'), (') − E', None), ('i,y', 'sub'), ('(0) ] / E', None),
         ('i,y', 'sub'), ('(0) × 100%', None)])
para('其中 θᵢ 为由 1981–2010 年历史风向频率在 0°–170° 范围内按 10° 步长扫描选定的'
     '气候加权最优角，G 在独立的 2014–2024 年逐时天气中回测；场级 Gᵢ 为所有可用'
     '年份的算术平均。选角与评价时期分离，模拟“建设前用历史气候选角、在未来天气中'
     '兑现收益”的决策过程；θᵢ 在评价期前冻结，个别年份 G 可以为负，这反映天气抽样'
     '而非重新优化失败。', size=9, indent=0.74, space_after=8)
para('公式 (S2)：风速年际贡献的标准差 SD(MS)——各场自然风速波动尺度', size=10.5,
     bold=True, space_after=2)
eq_para([('σ', None), ('S,i', 'sub'), (' = [ (1/(Y', None), ('i', 'sub'),
         ('−1)) Σ', None), ('y', 'sub'), (' (MS', None), ('i,y', 'sub'),
         (' − MS̄', None), ('i', 'sub'), (')² ]', None), ('1/2', 'sup')])
para('MSᵢ,ᵧ 为“实际/历史风速 × 实际/历史风向”四情景下年发电量偏差经二因素 Shapley '
     '分解分配的风速贡献。使用标准差而非相对固定历史基线的均方根，因为后者包含与'
     '年份无关的持续偏移，不是纯粹的年际变化。', size=9, indent=0.74, space_after=8)
para('公式 (S3)：场址条件化方向响应幅度 A——同一排布对朝向的敏感程度', size=10.5,
     bold=True, space_after=2)
eq_para([('A', 'i'), ('i', 'sub'), (' = [ max', None), ('θ', 'sub'),
         (' E', None), ('clim', 'sub'), ('(θ) − ⟨E', None), ('clim', 'sub'),
         ('(θ)⟩', None), ('θ', 'sub'), (' ] / ⟨E', None), ('clim', 'sub'),
         ('(θ)⟩', None), ('θ', 'sub'), (' × 100%', None)])
para('使用 1981–2010 年逐日风速—风向联合分布，在 0°–170° 范围内按 10° 步长扫描刚性'
     '旋转响应，取所有扫描相位中的最大气候加权年发电量相对旋转平均值的差。A 与 G '
     '不同：前者是历史气候与统一机型下终期排布的方向响应幅度（场址朝向潜力），后者是'
     '历史期选定角度在独立天气中的样本外增益。', size=9, indent=0.74, space_after=8)
para('公式 (S4)：圆周谐波幅度 Rₙ——风向玫瑰的轴向集中度', size=10.5, bold=True,
     space_after=2)
eq_para([('R', None), ('n', 'sub'), (' = | Σ', None), ('θ', 'sub'),
         (' w(θ) e', None), ('inθ', 'sup'), (' |', None)])
para('其中 w(θ) 为 5° 扇区的方向权重，分别取频率权重与能量权重两种口径；R1 即常用'
     '的风向集中度，R2 为二阶轴向集中度。由于同型机组在风向反转约 180° 时近似交换'
     '上下游位置，场级方向尾流函数近似满足 L(θ) = L(θ + 180°)，其响应主要由二阶及'
     '四阶谐波承载，因此 R2 在物理上是与阵列响应同阶的匹配量，而 R1 会使相反的季风'
     '分量相互抵消。', size=9, indent=0.74, space_after=8)
para('口径注：全文严格区分三类量——G 为建成朝向到最优朝向的模型化增益，A 为终期'
     '排布在给定长期风况与功率曲线加权下的方向响应幅度，ΔE 为项目情景下的额外年'
     '发电量。固定风速方向端元、长期边际分布诊断和逐时联合气候回测之间不作同口径'
     '数值比较。', size=9, indent=0.74)

h2('S1.4  统计检验与软件')
para('除特别说明外，所有相关采用双侧 Spearman 秩相关；阈值判别使用受试者工作特征'
     '曲线下面积（AUC）。记录长度主分析要求至少 5 年，3 年阈值仅作敏感性检验。'
     '空间模型的随机过程固定种子，地图全部采用海洋掩膜。计算环境为 Python 3 与 '
     'FLORIS v4.6（含 gauss/crespo_hernandez/sosfs 湍流模型同口径对照）。',
     indent=0.74)
doc.add_page_break()

# ══════════════ S2 长尾与越线 ══════════════
h1('S2  朝向敏感性长尾与越线风场（支撑正文 2.1）')
para('全样本 1,203 个风场—年中 64.3% 的增益为正，但 171 个风场的多年平均增益中位数'
     '仅为 0.34%、均值为 0.95%——全球均值主要由少数高响应场址抬高。表 S1 列出满足 '
     'G > SD(MS) 的逐场越线项目。主判据要求至少 5 个完整年份'
     '（108 个风场中 5 个越线，4.6%）；至少 3 年样本作为记录长度敏感性检验'
     '（146 个风场中 6 个越线）。该比较只表明这些项目的可控朝向收益超过其自身风速'
     '年际波动，并不意味着朝向在全球范围内比平均风速更重要。', indent=0.74)

table_caption('表 S1｜逐场越线项目（G > SD(MS)）')
add_table(
    ['场址编号', '海域 / 走廊', '主稿证据（逐年模型化增益 G）'],
    [
        ['F57', '越南季风海岸', '11 年中机组数由 48 台增至 80 台，逐年 G 保持在 16%–22%'],
        ['F66', '杭州湾口（东海）', '尾流端元分解见表 S2'],
        ['F91', '珠江口（南海）', '同走廊高增益项目'],
        ['F155', '塔兰托湾（地中海）', '仅在 ≥3 年样本中越线'],
        ['F157', '丹麦海峡', '机组数基本不变，逐年 G 约 5.0%–9.6%'],
        ['F160', '美国东海岸', 'G/SD(MS) = 1.03，主判据样本中比值最低、恰过 1:1 线'],
    ],
    col_widths=[2.4, 4.4, 9.2],
    note='注：六场比值（G / SD(MS)）范围 1.03–2.85（主稿正文 2.1）。主判据（≥5 年，'
         '108 场）下 5 场越线（4.6%）：F57、F66、F91、F157、F160；F155 仅在放宽到 '
         '≥3 年（146 场）样本中越线。逐场比值与逐年序列存于冻结派生数据，将随数据'
         '可用性声明一并公开。同一走廊中的项目仍可因建成相位和风速年际波动不同而'
         '不越线，走廊背景不能替代逐场比较。')

# ══════════════ S3 机制证据 ══════════════
h1('S3  高朝向响应的机制证据（支撑正文 2.2）')
h2('S3.1  F66 尾流端元（表 S2）')
table_caption('表 S2｜F66 终期真实机位在 9 m s⁻¹ 下的尾流端元')
add_table(
    ['方向', '场级效率 η', '全场最小轮毂风速'],
    [
        ['θ = 0°（高尾流方向）', '14.6%', '3.42 m s⁻¹'],
        ['θ = 35°（低尾流方向）', '91.8%', '8.27 m s⁻¹'],
    ],
    col_widths=[5.6, 4.4, 6.0],
    note='注：两个端元之间机组数量与两两距离保持不变。跨全部特征完整风场，第 90 百分位'
         '速度亏损的方向差与 A 的 Spearman 相关为 0.790（p = 5.8×10⁻³⁷）；建成朝向下的'
         '平均尾流损失与 A 几乎不相关（ρ = −0.045）。总尾流损失与尾流对方向变化的敏感'
         '程度因此是不同的量。')

h2('S3.2  全样本相关与偏相关（表 S3）')
table_caption('表 S3｜171 场全样本 Spearman 相关（A 两口径：真实排布 A_real 与六范式均值 A_para）')
add_table(
    ['变量', 'A_real（ρ, p）', 'A_para 六范式均值（ρ, p）'],
    [
        ['有效最近邻间距 S（D）', '−0.749（4.8×10⁻³²）', '−0.314（2.8×10⁻⁵）'],
        ['一阶集中度 R1_f（频率）', '+0.057（0.46）', '+0.045（0.56）'],
        ['二阶轴向集中度 R2_f（频率）', '+0.610（8.8×10⁻¹⁹）', '+0.665（3.6×10⁻²³）'],
        ['一阶集中度 R1_e（能量）', '−0.051（0.51）', '−0.018（0.82）'],
        ['二阶轴向集中度 R2_e（能量）', '+0.584（5.1×10⁻¹⁷）', '+0.665（3.1×10⁻²³）'],
        ['平均风速 ws', '−0.572（3.0×10⁻¹⁶）', '−0.630（2.8×10⁻²⁰）'],
        ['排布分散度 disp_f', '−0.057（0.46）', '−0.045（0.56）'],
    ],
    col_widths=[6.4, 5.0, 5.6],
    note='秩偏相关（控制第三方）：ρ(A, R2_e | S) = +0.585（4.7×10⁻¹⁷）；'
         'ρ(A, S | R2_e) = −0.732（6.5×10⁻³⁰）；ρ(A, R1_e | S) = −0.064（0.40）；'
         'ρ(A, S | ws) = −0.707（3.1×10⁻²⁷）。'
         '对数 A 联合回归（OLS R²，仅作叙事性比较）：仅 S 0.561；仅 R2_e 0.341；'
         'S + R2_e 0.693；S + R2_e + ws 0.775。')

h2('S3.3  受控气候反事实 C0–C3（表 S4）')
table_caption('表 S4｜受控气候反事实下的 A（171 个场址 × 六种标准化排布平均）')
add_table(
    ['情景', '风况设定', 'A', '相对 C0 变化'],
    [
        ['C0', '观测风速—风向联合分布', '1.09%', '—'],
        ['C1', '风速与风向解耦（保留各自边际）', '0.75%', '−0.34 pp（−30.9%）'],
        ['C2', '在 C1 基础上统一风速边际', '0.72%', '−0.37 pp'],
        ['C3', '仅把风向改为均匀分布', '0', '−1.09 pp'],
    ],
    col_widths=[1.6, 8.2, 2.2, 4.0],
    note='注：C1 中非均匀风向仍被保留，但风速—风向耦合被去除；C3 使风向均匀化后 A 严格'
         '为 0，即非均匀风向是产生模型化朝向响应的必要条件。风速—风向耦合贡献约三成，'
         '统一风速边际带来的附加变化较小。该比较衡量的是受控反事实中的敏感性，不是运行'
         '观测中的独立因果效应。')

h2('S3.4  分层判别（表 S5）')
table_caption('表 S5｜高响应分层判别（A_real > 5.2% 为高响应；密排 = S < 4D；高轴向 = R2_e > 中位）')
add_table(
    ['分层', 'n', '高响应场数', '占比'],
    [
        ['全部', '171', '27', '15.8%'],
        ['密排（S < 4D）', '123', '27', '22.0%'],
        ['　密排 × 高轴向', '67', '22', '32.8%'],
        ['　密排 × 低轴向', '56', '5', '8.9%'],
        ['疏排（S ≥ 4D）', '48', '0', '0.0%'],
        ['　疏排 × 高轴向', '18', '0', '0.0%'],
        ['　疏排 × 低轴向', '30', '0', '0.0%'],
    ],
    col_widths=[5.4, 2.6, 4.0, 4.0],
    note='注：小间距是几何门槛（必要条件）——27 个高响应风场 100% 在密排内、48 个疏排'
         '风场无一越线；轴向风气候在门槛内塑造幅度——密排 × 高轴向 32.8% vs 密排 × '
         '低轴向 8.9%（3.7 倍判别）。窄风玫瑰（一阶集中度 R1）不是高响应的可靠解释：'
         '区分力接近零、信号跨风数据不稳定；可靠的气候侧变量是二阶轴向集中度 R2。')
h2('S3.5  二阶谐波相位重构（主稿 5.6）')
para('阵列刚性旋转的主方向响应近似具有 180° 周期。将响应幅度 A 与历史最优相位组合的'
     '二阶谐波重构，与 2024 年逐时模拟增益比较：Spearman 相关 0.773、Pearson 相关 '
     '0.776，平均绝对误差 0.96 个百分点、平均偏差 +0.66 个百分点；A 单独的相关仅为 '
     '0.449，相位差单独为 0.474。约 13% 的风场在 2024 年出现负增益——响应幅度决定'
     '潜力，建成相位及跨期风况共同决定能否兑现。27 个高响应风场中，建成—最优相位差'
     '与 100G/A 的 Spearman 相关为 0.861（主稿图 2d）。该分析是二阶结构重构而非独立'
     '训练—测试预测，重构只用于检验幅度与相位能否共同恢复主要排序。', indent=0.74)
doc.add_page_break()

# ══════════════ S4 建设范式 ══════════════
h1('S4  建设范式与全圆最优（支撑正文 2.3）')
h2('S4.1  六套标准化建设范式的定义（表 S6、图 S1）')
table_caption('表 S6｜六套标准化建设范式的定义（每套 64 台 IEA 10 MW，D = 198 m）')
add_table(
    ['范式', '名称', '原型匹配场数', '生成规则'],
    [
        ['S_A', '侧风对齐', '41', '行轴垂直于能量加权风向'],
        ['S_B0', '约束优先（正北轴）', '36（S_B0/S_B45 合计）', '固定地理轴正北 0°'],
        ['S_B45', '约束优先（45° 轴）', '同上', '固定地理轴 45°'],
        ['S_C', '分期扩建', '100', '一期核心区 + 沿侧风向扩建'],
        ['S_D', '风资源梯度', '24', '按风功率密度梯度优先选点'],
        ['S_E', '大间距', '9', 'S_A 布局间距 × 1.25'],
    ],
    col_widths=[1.8, 3.6, 4.2, 6.4],
    note='注：原型匹配场数为与各范式形态—间距类别匹配的真实风场数；情境层排布为在'
         '原型场址上按规则生成的 64 机位标准模板，最小间距数值见表 S7。')

doc.add_picture(FIG_S1, width=Cm(15.0))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
caption('图 S1｜六套标准化建设范式的 64 机位排布示意图（面板 a–f 对应表 S6 的 S_A–S_E；'
        '原型场址；每套 64 台 IEA 10 MW，D = 198 m）。红色箭头为风来流方向。图内文字为'
        '英文，字体与主稿图 1–4 一致。该图为排布形态示意，情境层最小间距见表 S7，与'
        '原型场址间距属不同口径。')

h2('S4.2  六范式全圆最优（表 S7）')
table_caption('表 S7｜六套范式在 0°–170° 全圆扫描下的最优结果（相对建成朝向 0° 基线）')
add_table(
    ['范式', '最小间距', '全圆最优 ΔE 中位', '[5–95%]', 'ΔE 总量（GWh yr⁻¹）', 'GWh/GW（走廊 | 其他）'],
    [
        ['S_A 侧风对齐', '9.44D', '+9.62%', '[+0.4, +28.8]', '+70,284', '346 | 482'],
        ['S_B0 约束轴 0°', '3.33D', '−11.70%', '—', '−52,801', '−362 | —'],
        ['S_B45 约束轴 45°', '3.33D', '−11.70%', '—', '−52,802', '−353 | —'],
        ['S_C 分期扩建', '5.00D', '+0.98%', '—', '+21,081', '58 | 148'],
        ['S_D 风资源梯度', '4.00D', '−5.63%', '—', '−17,758', '−155 | −116'],
        ['S_E 大间距', '11.80D', '+10.79%', '—', '+75,966', '376 | 520'],
    ],
    col_widths=[3.2, 1.9, 2.9, 2.6, 3.2, 3.2],
    note='注：走廊 23 场 vs 其他 148 场的 A 中位倍比为 5.5–7.4 倍（逐情景 p < 1×10⁻¹⁰）；'
         '六范式场址排序两两 Spearman 相关中位数为 0.96；走廊 A 中位依次为 1.2%、5.3%、'
         '5.3%、2.9%、4.6% 和 0.8%。密集排布的绝对响应最强，即使在 11.8D 大间距范式'
         '中，走廊响应仍为其他场址的 5.5 倍。')

h2('S4.3  机型与年代匹配敏感性（表 S8）')
table_caption('表 S8｜机型与年代匹配敏感性（走廊/其他倍比范围、场址排序一致性）')
add_table(
    ['机型口径', '走廊/其他倍比范围', '类型匹配 Spearman', '两两 Spearman 中位'],
    [
        ['nrel_5MW（D = 126 m）', '5.4–6.7×', '0.711', '0.958（min 0.913）'],
        ['iea_10MW（D = 198 m，基准）', '5.5–7.4×', '0.706', '0.961（min 0.912）'],
        ['iea_15MW（D = 242 m）', '5.6–7.4×', '0.709', '0.963（min 0.913）'],
        ['年代匹配（各场建设年代机型）', '5.1–6.8×', '0.718', '0.955'],
    ],
    col_widths=[5.6, 3.8, 3.6, 4.0],
    note='注：类型匹配 Spearman 为各机型口径下范式均值 A 与真实排布 A（iea_10MW 口径）'
         '的相关。能量加权风向 θ_energy 与基准的差：iea_10MW 中位 0.00°（P90 0.0°）、'
         'iea_15MW 中位 0.12°（P90 0.7°）。年代匹配口径机型分布为 nrel_5MW 99 场、'
         'iea_10MW 44 场、iea_15MW 28 场（未知年份按 iea_10MW）。')

h2('S4.4  气候与范式的方差分解（表 S9）')
table_caption('表 S9｜六情景响应的方差分解（171 场 × 6 情境，η²）')
add_table(
    ['因素', '解释方差份额'],
    [
        ['场址气候主效应（风速与风向）', '60.72%'],
        ['建设范式主效应', '19.27%'],
        ['气候 × 范式交互', '20.01%'],
    ],
    col_widths=[9.6, 4.4],
    note='注：包含风速与风向的场址气候主效应解释 60.7% 的响应方差，但该项为混合气候'
         '项（含风速与风向的联合作用），不应被单独解释为风向因果贡献。均匀风向负对照'
         '使 A 降为 0（表 S4 C3），说明非均匀风向组织是形成模型化朝向响应的必要条件。')

# ══════════════ S5 走廊与反事实 ══════════════
h1('S5  方向性走廊与项目级平行反事实（支撑正文 2.4）')
h2('S5.1  空间筛查模型的整组留出（表 S10）')
table_caption('表 S10｜整组留出验证（物理走廊整组留出 / 国家整组留出）')
add_table(
    ['评价目标', 'AUC', 'Spearman 秩相关', '前四分位捕获率'],
    [
        ['范式均值 A（简化筛查模型）', '0.628', '0.415', '35.7%'],
        ['真实排布 A（同一模型）', '0.498', '0.087', '31.0%'],
    ],
    col_widths=[7.6, 2.4, 3.4, 4.0],
    note='注：统计模型预测值相对两种标签的 Spearman 相关为 0.680（范式均值 A）与 '
         '0.609（真实 A）。管线设计目标（AUC ≥ 0.75、前 25% 捕获 ≥ 50%）未达到——'
         '整走廊留出仅用于界定简化筛查模型的外推边界，不评价六范式或真实排布直接复算'
         '的准确性；候选地图不能替代未知项目的直接方向响应计算。'
         'C3 均匀风向负对照：农场与格点 max|A| = 0 pp（通过 ≤ 0.1 pp 验收红线）。'
         '角度分辨率敏感性（5°→1°）：|ΔA| p50 = 0.005、p95 = 0.058 pp'
         '（通过 ≤ 0.3 pp 红线）。')

h2('S5.2  项目级平行反事实（表 S11）')
para('项目级分析纳入机组数不少于 10 的 155 个已建项目，以统一 IEA 10 MW 机型计算并按'
     '项目机组数换算容量，总覆盖 149.2 GW。三类情景均以建成排布在 0° 基线的年发电量'
     '为参照，使用 2014–2019 年选择最优朝向，并在 2020–2024 年逐年样本外评价。',
     indent=0.74)
table_caption('表 S11｜S1/S2/S3 平行反事实：走廊（20 项目，11.28 GW = 7.6% 装机）与其他海域对比')
add_table(
    ['情景', '全样本中位', '走廊中位', '其他中位', '走廊 [5–95%]', '其他 [5–95%]',
     '走廊 ΔE（GWh，占全样本）', 'GWh/GW 走廊 vs 其他'],
    [
        ['S1 仅校正朝向', '+0.5%', '+2.77%', '+0.40%', '[+0.61, +7.45]',
         '[+0.00, +2.19]', '995（28.6%）', '88 vs 18'],
        ['S2 匹配模板重排', '+1.6%', '−1.36%', '+2.28%', '[−27.02, +12.25]',
         '[−16.59, +22.38]', '−626（−3.9%）', '−56 vs 120'],
        ['S3 六模板择优', '+10.8%', '+10.37%', '+10.86%', '[+1.51, +25.67]',
         '[+1.46, +29.81]', '4,244（5.6%）', '376 vs 520'],
    ],
    col_widths=[2.8, 1.8, 1.8, 1.8, 2.2, 2.2, 2.4, 2.4],
    font_size=8.5,
    note='注：全样本 ΔE 为 S1 3,474 / S2 15,889 / S3 75,966 GWh yr⁻¹。S3 下增益超过 '
         '5% 的项目占比走廊 80.0% vs 其他 88.1%，超过 10% 为 60.0% vs 56.3%。'
         'S2 与 S3 是相对同一建成基线的平行、非嵌套反事实，其候选集中不必包含真实排布'
         '或 S1 解，因此相对建成基线可以为负；这些结果称为发电变化，而不预设为可避免'
         '损失。S3 未显式施加租区面积、海缆、排除区和成本约束，解释为标准模板库内的'
         '无约束上界。')

h2('S5.3  分走廊区域结果（表 S12）')
table_caption('表 S12｜分走廊区域结果（n ≥ 10 口径，155 项目、149.2 GW）')
add_table(
    ['走廊 / 区域', '项目数', '装机（GW）', 'S1 中位', 'S2 中位', 'S3 中位',
     'ΔE S1（GWh）', 'ΔE S2（GWh）', 'ΔE S3（GWh）'],
    [
        ['台湾海峡', '8', '6.13', '+2.92%', '+2.47%', '+15.36%', '659.3', '216.7', '2,714.4'],
        ['越南', '12', '5.15', '+1.58%', '−5.66%', '+8.30%', '335.9', '−843.2', '1,529.4'],
        ['其他海域', '135', '137.87', '+0.40%', '+2.28%', '+10.86%', '2,478.9', '16,515.3', '71,722.4'],
    ],
    col_widths=[2.6, 1.6, 2.0, 1.7, 1.7, 1.7, 2.0, 2.0, 2.0],
    font_size=8.5,
    note='注：仅校正朝向的 S1 收益明显集中于走廊（28.6% 的新增电量来自 7.6% 的装机）；'
         '标准模板替换（S2）和六模板择优（S3）同时改变间距与几何，结果并不在走廊内'
         '集中。规划含义因而是有选择地保留方向自由度，并在具体租区和工程约束下复核，'
         '而不是对全球海上风电施加统一的朝向增益率。')

# ══════════════ S6 敏感性分析可视化 ══════════════
h1('S6  敏感性分析可视化')
para('图 S3 汇总四个维度的敏感性证据。(a) 记录长度敏感性：将主判据从至少 5 年放宽到'
     '至少 3 年，越线判定由 5/108（4.6%）变为 6/146（4.1%），两组 G 的累积分布几乎'
     '重合，结论对记录长度阈值不敏感。(b) 机型口径敏感性（表 S8）：走廊/其他响应'
     '倍比在 NREL 5 MW、IEA 10 MW（基准）、IEA 15 MW 与年代匹配口径下均保持 5.1–7.4'
     ' 倍，机型口径不改变走廊强化的结论。(c) 受控气候反事实（表 S4）：风速—风向'
     '解耦使 A 下降约三成，风向均匀化使 A 严格为 0——非均匀风向结构是产生模型化'
     '朝向响应的必要条件。(d) 风速窗口机制分解（5.8 节补算，6 个越线场）：最优相位'
     '相对平均风向的可恢复能量以 7–10 m s⁻¹ 部分负荷窗口为主（57%–69%），15–25 '
     'm s⁻¹ 额定区贡献约 0，即功率曲线在部分负荷区完成能量转化；该诊断独立组合'
     '风向频率与风速边际，不与逐时联合气象的 G 作同口径数值比较。', indent=0.74)
doc.add_picture(FIG_S3, width=Cm(15.0))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
caption('图 S3｜敏感性分析可视化（图内文字为英文，字体与主稿图 1–4 一致）。(a) G 累积'
        '分布：n ≥ 5 主判据（n = 108）与 3 ≤ n < 5 附加样本（n = 38），红星为 6 个越线'
        '场。(b) 机型口径敏感性：走廊/其他 A 倍比的六范式区间（表 S8）。(c) 受控气候'
        '反事实 C0–C3 的 A（表 S4）。(d) 6 个越线场最优相位相对平均风向的可恢复能量'
        '按风速窗口分解（5.8 节补算，口径 B），场后标注可恢复百分比。')

# ══════════════ S7 统计口径与诚实边界 ══════════════
h1('S7  统计口径与诚实边界')
for t in [
    '高响应场数：正文统一使用 27/171（iea_10MW 统一机型 + 逐场湍流强度，当前冻结管线）。'
    '另一代口径 29/167 出自统一 nrel_5MW 的旧管线；两代 A 在共同 167 场上的 Spearman '
    '相关为 0.983，差异只发生在阈值边缘的 2 个场（F90 5.56→4.81、F145 6.75→4.83），'
    '管线升级未改变任何结论。',
    '走廊口径：物理走廊 23 场（越南季风海岸、中国东南湾口、亚得里亚海—塔兰托湾、'
    '丹麦海峡）；项目级分析因机组数 ≥ 10 要求对应 20 个项目。留出验证按整走廊或整'
    '国家成组排除，避免同走廊样本泄漏。',
    '相位几何代理：排布主轴角仅用于描述总体对齐，其边际分布不证明排布方向与风况相互'
    '独立，更不用于识别岸线、航道、海缆或租区约束的因果作用。',
    '整走廊留出以真实排布 A 为目标时 AUC = 0.498，接近随机：简化筛查模型的外推能力'
    '有限，正文据此将其限定为外推边界界定工具，而非项目级决策依据。',
    '受控气候反事实（C0–C3）与固定风速端元均为模型内诊断，衡量受控仿真中的敏感性，'
    '不是运行观测中的独立因果效应。',
    '走廊在 S2 中的中位发电变化为负（−1.36%）：建成排布优于范式标准化重排，该结果'
    '如实报告为“发电变化”而非“可避免损失”；走廊剩余空间已主要经 S1 兑现（28.6%）。',
    '数据与代码：支撑各图的场级与风场—年派生数据、方向响应曲线、整组留出预测及绘图'
    '脚本将在论文接收时存入具有持久标识符的公共仓库，并提供从正文数字到冻结派生数据'
    '和计算步骤的溯源表（与主稿“数据可用性”“代码可用性”一致）。',
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.74)
    r = p.add_run('• ' + t)
    set_font(r, size=10.5)

# ══════════════ 补充参考文献 ══════════════
h1('补充参考文献（Supplementary References）')
para('编号 S1–S4 与主稿参考文献编号对应，完整著录见主稿参考文献列表；S5、S6 为本'
     '补充信息新增。', size=9)
refs = [
    '[S1] 主稿参考文献[14]：全球公开海上风机机位数据集。Zhang T, Tian B, Sengupta D, '
    'Zhang L, Si Y. Global offshore wind turbine dataset. Sci Data. 2021;8:191. '
    'https://doi.org/10.1038/s41597-021-00982-z',
    '[S2] 主稿参考文献[15]：ERA5 全球再分析。Hersbach H, Bell B, Berrisford P, et al. '
    'The ERA5 global reanalysis. Q J R Meteorol Soc. 2020;146:1999–2049. '
    'https://doi.org/10.1002/qj.3803',
    '[S3] 主稿参考文献[16]：再分析风模拟的区域偏差。Gruber K, Regner P, Wehrle S, '
    'Zeyringer M, Schmidt J. Towards global validation of wind power simulations: '
    'A multi-country assessment of wind power simulation from MERRA-2 and ERA-5 '
    'reanalyses bias-corrected with the Global Wind Atlas. Energy. 2022;238:121520. '
    'https://doi.org/10.1016/j.energy.2021.121520',
    '[S4] 主稿参考文献[17]：高斯尾流模型。Bastankhah M, Porté-Agel F. A new '
    'analytical model for wind-turbine wakes. Renew Energy. 2014;70:116–123. '
    'https://doi.org/10.1016/j.renene.2014.01.002',
    '[S5] NREL. FLORIS Version 4.6 [Computer software]. Golden, CO: National '
    'Renewable Energy Laboratory. https://github.com/NREL/floris',
    '[S6] Bortolotti P, Tarres HC, Dykes K, et al. IEA Wind TCP Task 37: Systems '
    'Engineering in Wind Energy – WP2.1 Reference Wind Turbines. Technical Report '
    'NREL/TP-5000-73492. Golden, CO: National Renewable Energy Laboratory; 2019.',
]
for t in refs:
    para(t, size=9, space_after=4)

doc.save(OUT)
print('已生成:', OUT)
