"""
论文初稿回填：Nature Energy 中文稿（幅度相位框架 v3.0）→ v3.1
================================================================
按方案 §4.9/§5.9 回填模板 [77][79][118] 与 [135]（更新图3、绘制图4，
数字回填 Result 3/4、讨论、结论三/四），全部数字从计算输出动态提取：

  结论三：wp5_cross_farms/grid.npz、wp5_anova.txt、wp5_report.txt、wp5_rfd_grid.csv
          wp6a_loo.npz、wp7a_real_curves.npz（存在则用自算 A，否则 task3）
  结论四：wp7b_scenario_table.csv、wp7b_corridor_summary.csv

改动段落（v3.0 索引）：
  [4] 稿件状态 | [6] 摘要 Result 4 回填 + 交叉仿真证据
  [12] 问题清单末句 | [40] Result 3 叙述 + 主证据链
  [42] 结论三（回填模板 [77]+[79]）| [44] 图3 图注（新五面板）
  [45] 2.4 标题 | [46] 图4 图注 + 插入图片 | [47] 主结论句回填
  [53] 讨论第四点 | [58] 结论第三点 + 交叉仿真 | [59] 结论第四点 + S1-S3
  [60] Result 4 回填位 | [82] 方法限制更新（网格计算已用 ERA5 本地档案完成）
  [98] 5.10 标题 | [104] 5.10 末追加完成状态
  [43] 图3 图片替换为新版五面板 | [46] 前插入图4 图片

输出：四场景风速风向分解贡献/Nature_Energy_中文稿_幅度相位框架_v3.1_结论三四回填.docx
"""
import os, io, sys, warnings, shutil, re, glob
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
warnings.filterwarnings('ignore')
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from scipy.stats import spearmanr

BUSH = os.path.dirname(os.path.abspath(__file__))
REPO = os.path.dirname(BUSH)
OUT = os.path.join(BUSH, 'output')
SRC_DIR = os.path.join(REPO, '四场景风速风向分解贡献')
SRC = glob.glob(os.path.join(SRC_DIR, 'D_onedrive01_科研与论文08_风向建设分析材料_朝向长尾走廊Nature_Energy_中文稿_幅度相位框架_v3.0 2(1).docx'))[0]
DST = os.path.join(SRC_DIR, 'Nature_Energy_中文稿_幅度相位框架_v3.1_结论三四回填.docx')

# ═══════════════════════════════════════════════════════════════════════
# 0. 数字提取
# ═══════════════════════════════════════════════════════════════════════
z5 = np.load(os.path.join(OUT, 'wp5_cross_farms.npz'))
A_farm = z5['A']; farm_ids = z5['farm_ids'].astype(int)
z5g = np.load(os.path.join(OUT, 'wp5_cross_grid.npz'))
rfd = pd.read_csv(os.path.join(OUT, 'wp5_rfd_grid.csv'), encoding='utf-8-sig')
gv = z5g['valid']
# ANOVA η²
anova = open(os.path.join(OUT, 'wp5_anova.txt'), encoding='utf-8').read()
eta2 = {m.strip(): float(v) for m, v in re.findall(r'(气候\(场址\)|形态\s*|间距\s*|场址×形态|场址×间距|形态×间距|残差)\s+SS=[\d.]+ +η²=([\d.]+)%', anova)}
# Jaccard / R-Spearman
rep5 = open(os.path.join(OUT, 'wp5_report.txt'), encoding='utf-8').read()
jaccard = float(re.search(r'Jaccard 均值: ([\d.]+)', rep5).group(1))
r_rmap = float(re.search(r'Spearman: ([\d.]+)', rep5).group(1))
# F75 格点比例（CSV 为 0-100 百分比刻度：k/36 模板 × 100）
F = rfd['F75'].values[gv]
f50 = float((F >= 50).mean() * 100)
f75 = float((F >= 75).mean() * 100)
# C0-C3
A_c = z5['A_c']
c0, c3 = np.nanmean(A_c[:, :, 0]), np.nanmean(A_c[:, :, 3])
# 类型匹配（wp7a 自算 A 优先）
z7p = os.path.join(OUT, 'wp7a_real_curves.npz')
if os.path.exists(z7p):
    z7 = np.load(z7p)
    A_real_s = pd.Series(z7['A'], index=z7['farm_ids'].astype(int))
    real_src = 'WP7a 自算 36 档半圆口径'
else:
    t3 = pd.read_csv(os.path.join(REPO, 'task3', 'task3_s1_optimal_orientation.csv'), encoding='utf-8-sig')
    g3 = t3.groupby('farm_id')['expected_AEP_kWh'].agg(['max', 'mean'])
    A_real_s = 100 * (g3['max'] - g3['mean']) / g3['mean']
    real_src = 'task3 18 角度口径'
common = [i for i, f in enumerate(farm_ids) if f in A_real_s.index]
rho_type = spearmanr(A_farm[common].mean(axis=1), A_real_s.loc[farm_ids[common]].values)[0]
# 嵌套留出
zloo = np.load(os.path.join(OUT, 'wp6a_loo.npz'))
_w6a = open(os.path.join(OUT, 'wp6a_report.txt'), encoding='utf-8').read()
auc_loo = float(re.search(r'AUC=([\d.]+)', _w6a).group(1))   # 首个 AUC = 物理模板A均值标签
# 结论四
tb = pd.read_csv(os.path.join(OUT, 'wp7b_scenario_table.csv'), encoding='utf-8-sig')
main = tb[tb.layout_type != 'sparse']
s1, s2, s3 = main.G_plan_S1.median(), main.G_plan_S2.median(), main.G_plan_S3.median()
gwh = main[['dE_GWh_S1', 'dE_GWh_S2', 'dE_GWh_S3']].sum()
cap_gw = main.capacity_MW.sum() / 1000
v1 = main.V1.sum() / main.V1.abs().sum()
v2 = main.V2.sum() / main.V1.abs().sum()
v3 = main.V3.sum() / main.V1.abs().sum()
mono = main.monotonic.mean() * 100
pos_s1 = main.pos_frac_S1.median() * 100

# ═══════════════════════════════════════════════════════════════════════
# 1. 载入初稿并复制
# ═══════════════════════════════════════════════════════════════════════
shutil.copy(SRC, DST)
doc = Document(DST)
paras = doc.paragraphs

def set_text(p, text):
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    p.add_run(text)

P = lambda i: paras[i]

# ═══════════════════════════════════════════════════════════════════════
# 2. 文本回填
# ═══════════════════════════════════════════════════════════════════════
set_text(P(4), '稿件状态｜Result 1–3 已按“场址条件化响应幅度—项目相位兑现”统一，并完成受控交叉仿真与'
               '嵌套整走廊留出验证；Result 4 已按补算方案降级路径完成已建风场的三类建设情境方法验证，'
               '全球规划容量层保留项目级开发与工程约束情景的定量回填位。')

t6 = P(6).text
t6 = t6.replace('〔Result 4待补：纳入核算的项目容量、不同开发与工程情景的新增电量及机会保留率。〕',
    f'受控交叉仿真将响应幅度方差分解为气候主效应{eta2["气候(场址)"]:.1f}%、间距{eta2["间距"]:.1f}%与形态{eta2["形态"]:.1f}%，'
    f'均匀方向负对照下响应为零；高响应海域在36套标准排布中的位置高度一致（前四分位Jaccard {jaccard:.2f}）。'
    f'对已建风场的三类建设情境中，仅朝向校正的中位增益为+{s1:.1f}%，保持排布类型的标准化重排低于建成基线（{s2:.1f}%），'
    f'联合重构排布形态、间距与朝向的中位增益为+{s3:.1f}%；净新增{gwh.sum():.0f} GWh yr⁻¹（覆盖{cap_gw:.0f} GW），'
    f'结果仅代表已覆盖项目。')
set_text(P(6), t6)

t12 = P(12).text
t12 = t12.replace('本稿完成前三项结果，并为第四项结果保留可直接回填的结构。',
                  '本稿完成前三项结果，并在已建风场样本上完成第四项结果的三类建设情境方法验证；全球规划容量层保留可直接回填的结构。')
set_text(P(12), t12)

set_text(P(40), P(40).text + f'受控交叉仿真（{int(gv.sum())}个有效海洋格点×36套标准排布）将响应幅度方差分解为'
    f'气候主效应{eta2["气候(场址)"]:.1f}%、间距{eta2["间距"]:.1f}%、场址×间距{eta2["场址×间距"]:.1f}%与形态{eta2["形态"]:.1f}%，'
    f'均匀方向负对照的响应为零（农场最大 {np.nanmax(np.abs(A_c[:, :, 3])):.6f} pp）；高响应格点在36套排布中的位置高度一致'
    f'（前四分位Jaccard {jaccard:.2f}）。这使图3的稳健排序地图由统计外推旁证升级为受控仿真主证据：'
    f'走廊首先是场址风速—风向气候的稳定属性，建设排布主要调节幅度而非位置。')

set_text(P(42), f'结论 3｜受控交叉仿真表明，方向性走廊首先是场址风速—风向气候的稳定属性，而非排布形态的产物。'
    f'在控制机型、容量与模板规模后，气候主效应解释响应幅度方差的{eta2["气候(场址)"]:.1f}%，间距与形态合计解释约'
    f'{eta2["间距"] + eta2["形态"]:.1f}%；均匀方向负对照下响应降为零。高响应海域在36套标准排布中的位置高度一致'
    f'（前四分位Jaccard {jaccard:.2f}），{f50:.0f}%的候选格点在至少18套、{f75:.0f}%在至少27套建设情景中保持高响应，'
    f'类型匹配验证达到Spearman {rho_type:.2f}；建设排布主要调节响应幅度，而不改变其主要空间位置。'
    f'朝向机会的全球性因此不是平均高收益的普遍性，而是方向性走廊形成机制能够跨区域重复。')

set_text(P(44), f'图3｜受控交叉仿真与稳健走廊。（a）171个真实风场×36套标准排布的响应幅度矩阵，'
    '行按场址幅度均值降序，列按形态×间距分组。（b）跨排布中位优先度分位R的全球观测支持域地图，'
    '星标为已知走廊成员风场。（c）跨模板前四分位支持率F75地图。（d）物理模板响应幅度与真实排布响应幅度的'
    f'类型匹配散点，按整走廊留出分组着色（Spearman ρ={rho_type:.2f}）。（e）嵌套整走廊留出下，'
    '按预测筛选比例对应的观测高响应（前四分位）召回曲线；设计目标为前25%筛选捕获≥50%。'
    '地图色标表示稳健排序或支持率，而非项目增益百分比；正文所称“全球”指观测支持域（北半球样本），非完整全球海域。')

set_text(P(45), '2.4 三类建设情境下的走廊增发电量（已建风场方法验证）')

set_text(P(46), '图4｜三类建设情境的增发电量。（a）候选走廊与纳入核算项目的容量地图，点大小正比于项目容量。'
    '（b）S1固定方案朝向校正、S2类型保持再设计与S3规划前置综合设计的项目级百分比增益中位数及5–95%区间。'
    '（c）各走廊新增GWh yr⁻¹贡献。（d）V1/V2/V3增量价值分解。图内全部使用英文。')

set_text(P(47), f'在171个已建成风场样本上完成三类建设情境的方法验证（统一机型iea_10MW、统一容量口径、'
    f'同一尾流查表与历史期选角—样本外回测）：相对于建成朝向基线，仅进行固定方案朝向校正（S1）的增益中位为'
    f'+{s1:.1f}%；在保持排布类型的条件下以标准化模板重排内部机位（S2）的中位增量为{s2:.1f}%'
    f'（建成排布的旋转最优在约{100 - mono:.0f}%的项目上优于其类型匹配的标准化模板，类型保持重排不构成普遍增益）；'
    f'联合选择排布形态、间距与朝向（S3）的中位增益达到+{s3:.1f}%。三类情境净模型化新增电量{gwh.sum():.0f} GWh yr⁻¹'
    f'（覆盖{cap_gw:.0f} GW；S1 +{gwh.iloc[0]:.0f}、S2 {gwh.iloc[1]:.0f}、S3 +{gwh.iloc[2]:.0f}），'
    f'朝向校正（V1）与规划前置重构（V3）为正，类型保持重排（V2）为负'
    f'（V1/V2/V3相对V1量值的份额为{v1:.0%}/{v2:.0%}/{v3:.0%}）；'
    f'三情境增量单调递增仅在{mono:.0f}%的项目上成立，样本外正增益年份比例中位{pos_s1:.0f}%。'
    f'该结果仅代表已覆盖项目；旧版以国家规划容量、统一容量因子和国家平均响应计算的225 GW/23.5 TWh yr⁻¹'
    f'只保留为审计对照，全球规划容量层待项目管线与租区多边形数据回填。')

t53 = P(53).text
t53 = t53.replace('第四，全球投影的幅度外推尚未通过项目级验证；租区边界、排除区、海缆、出口容量、邻场尾流和拟建机型都可能压缩理论朝向机会。',
                  f'第四，全球投影的幅度外推已通过已建风场的项目级验证（类型匹配Spearman {rho_type:.2f}，真实排布A为{real_src}），'
                  '但租区边界、排除区、海缆、出口容量、邻场尾流和拟建机型都可能压缩理论朝向机会，'
                  '且S1—S3核算仅覆盖已建项目，全球规划容量层仍需项目管线数据。')
set_text(P(53), t53)

t58 = P(58).text
t58 = t58.replace('方向性走廊因而是一类可以跨区域识别的物理空间类型，而不是若干彼此孤立的高值项目。',
    f'受控交叉仿真进一步表明走廊首先是气候属性：气候主效应解释响应幅度方差的{eta2["气候(场址)"]:.1f}%，'
    f'高响应海域在36套标准排布下位置高度一致（前四分位Jaccard {jaccard:.2f}），排布主要调节幅度而非位置。'
    '方向性走廊因而是一类可以跨区域识别的物理空间类型，而不是若干彼此孤立的高值项目。')
set_text(P(58), t58)

t59 = P(59).text
t59 = t59.replace('朝向的系统价值不是来自对全球海上风电容量施加统一增益率，而是来自在少数高敏感走廊中为具体项目保留并利用可实施的方向自由度。',
    f'在已建风场样本上，朝向校正（S1）的中位增益为+{s1:.1f}%，联合选择形态、间距与朝向的规划前置重构（S3）'
    f'为+{s3:.1f}%，而类型保持的标准化重排（S2）为{s2:.1f}%——'
    f'说明朝向自由度价值的一部分只能在规划阶段通过排布自由度兑现。'
    '朝向的系统价值不是来自对全球海上风电容量施加统一增益率，而是来自在少数高敏感走廊中为具体项目保留并利用可实施的方向自由度。')
set_text(P(59), t59)

set_text(P(60), f'Result 4内部回填位（不属于正式结论）：在171个已建风场（覆盖{cap_gw:.0f} GW）上，'
    f'S1朝向校正、S2类型保持重排与S3规划前置重构的增益中位分别为+{s1:.1f}%、{s2:.1f}%与+{s3:.1f}%，'
    f'净新增{gwh.sum():.0f} GWh yr⁻¹（S1 +{gwh.iloc[0]:.0f}/S2 {gwh.iloc[1]:.0f}/S3 +{gwh.iloc[2]:.0f}）；'
    f'V1/V2/V3相对V1量值份额为{v1:.0%}/{v2:.0%}/{v3:.0%}（V2为负：类型保持重排减损价值）。'
    f'结果仅代表已覆盖项目，全球规划容量层待项目管线数据回填。')

t82 = P(82).text
t82 = t82.replace('所依赖的原始气象数据在当前工作区不可用；同一限制也阻断Result 3的网格计算。',
    '该诊断当前依赖的扇区级风玫瑰仅覆盖7个走廊风场；'
    'Result 3的网格气候计算已改用本地ERA5四区域逐时档案完成（1,446个海洋格点中1,205个位于档案覆盖区内，'
    '其余241个在档案边界外，如实记录为数据边界）。')
set_text(P(82), t82)

set_text(P(98), '5.10 Result 4项目级能源情景（已建风场方法验证；全球规划容量层待项目管线数据）')

set_text(P(104), P(104).text + f'已按补算方案降级路径完成已建风场方法验证：统一机型iea_10MW与统一项目容量口径，'
    f'统一FLORIS尾流查表（18风速档×72风向），三类嵌套可行集S1⊆S2⊆S3，'
    f'历史期2014–2019选角、2020–2024逐年样本外评价；结果仅代表已覆盖项目。')

# ═══════════════════════════════════════════════════════════════════════
# 3. 图片：替换图3、插入图4
# ═══════════════════════════════════════════════════════════════════════
FIG3 = os.path.join(BUSH, 'figures', 'Fig3_conclusion3.png')
FIG4 = os.path.join(BUSH, 'figures', 'Fig4_conclusion4.png')
assert os.path.exists(FIG3) and os.path.exists(FIG4), '图3/图4 PNG 不存在，先运行 figures 脚本'

p43 = paras[43]
for r in list(p43.runs):
    r._r.getparent().remove(r._r)
p43.alignment = WD_ALIGN_PARAGRAPH.CENTER
p43.add_run().add_picture(FIG3, width=Cm(16.0))

p46 = paras[46]
np46 = p46.insert_paragraph_before()
np46.alignment = WD_ALIGN_PARAGRAPH.CENTER
np46.add_run().add_picture(FIG4, width=Cm(16.0))

doc.save(DST)
print(f'完成 → {DST}')
print(f'结论三: 气候 η²={eta2["气候(场址)"]:.1f}% | Jaccard={jaccard:.2f} | 类型匹配 ρ={rho_type:.2f} '
      f'({real_src}) | F75≥50%格点 {f50:.0f}% | 嵌套留出 AUC={auc_loo}')
print(f'结论四: S1 {s1:.2f}% / S2 {s2:.2f}% / S3 {s3:.2f}% | ΔE 合计 {gwh.sum():.0f} GWh/yr '
      f'(覆盖 {cap_gw:.0f} GW) | V1/V2/V3={v1:.2f}/{v2:.2f}/{v3:.2f} | 单调 {mono:.0f}%')
