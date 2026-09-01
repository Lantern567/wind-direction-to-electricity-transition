# -*- coding: utf-8 -*-
"""图 1 (d)(e) 优化版换入 v4.9 + 图注同步（2026-09-02）
====================================================================
与 8-30 的 swap_fig1_into_v49.py 的区别：
  1. 媒体项编号已变：9-01 版 v4.9 重存过，图 1 从 image9.png 变成
     image1.png（按 4342×5286 尺寸定位，不再硬编码旧编号）。
  2. 图注中越线场数（4/1/两幅/5 条）8-30 已改过，本次不再动；只补两处
     新增图元的说明：(d) 虚线＝仅 3–4 年记录、(e) 对角参考线与阴影。
  3. 图片字节按同尺寸替换，文档声明的显示范围不变、无比例失真。
"""
import os, sys, io, shutil, hashlib, zipfile
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from PIL import Image

ROOT = r'D:\onedrive\01_科研与论文\08_风向建设'
DOCX = os.path.join(ROOT, 'ADAPEN_manuscript_zh_v4.9.docx')
BAK = os.path.join(ROOT, 'ADAPEN_manuscript_zh_v4.9_图1旧版备份_20260902.docx')
PNG = os.path.join(ROOT, 'wind-direction-to-electricity-transition', '重算交付_20260829',
                   '图件', '构建脚本', 'Fig1_senior_nat_v2.png')
FIG1_SIZE = (4342, 5286)

# ── 0. 定位图 1 的媒体项（按尺寸，不依赖编号）──
z = zipfile.ZipFile(DOCX)
targets = []
for n in z.namelist():
    if n.startswith('word/media/'):
        try:
            if Image.open(io.BytesIO(z.read(n))).size == FIG1_SIZE:
                targets.append(n)
        except Exception:
            pass
z.close()
assert len(targets) == 1, '按 %s 定位到 %d 个媒体项：%s' % (FIG1_SIZE, len(targets), targets)
MEDIA = targets[0]
print('图 1 媒体项:', MEDIA)

shutil.copy(DOCX, BAK)
print('备份:', os.path.basename(BAK))


def para_replace(p, old, new, tag):
    texts = [r.text for r in p.runs]
    full = ''.join(texts)
    assert old in full, '%s: 未找到 %r' % (tag, old[:40])
    i0 = full.index(old); i1 = i0 + len(old)
    pos = 0; si = None
    for k, t in enumerate(texts):
        if pos <= i0 < pos + len(t):
            si = k; break
        pos += len(t)
    assert si is not None
    prefix = texts[si][:i0 - pos]
    pos_k = pos; last = si
    while pos_k < i1:
        pos_k += len(texts[last]); last += 1
    tail = i1 - (pos_k - len(texts[last - 1]))
    suffix = texts[last - 1][tail:] if tail > 0 else ''
    p.runs[si].text = prefix + new + suffix
    for j in range(si + 1, last):
        p.runs[j].text = ''
    print('  -', tag)


# ── 1. 图注补两处新增图元说明 ──
d = Document(DOCX)
cap = None
for p in d.paragraphs:
    if p.text.strip().startswith('图 1｜'):
        cap = p; break
assert cap is not None, '未找到图 1 图注'
assert '虚线折线' not in cap.text and '均匀对应' not in cap.text, '图注已同步过，勿重复运行'
print('图注改动:')
para_replace(cap,
             '灰带与灰线为全部风场的四分位距与中位数，纵轴为对称对数刻度。',
             '灰带与灰线为全部风场的四分位距与中位数，虚线折线为仅具 3–4 年记录的风场，'
             '纵轴为对称对数刻度（1% 以下为线性段）。', '(d) 补虚线与线性段说明')
para_replace(cap,
             '橙点标出增益最高的 5 个风场。',
             '橙点标出增益最高的 5 个风场，对角虚线为装机份额与新增电量份额均匀对应的'
             '参考线，阴影为曲线相对该参考线的偏离。', '(e) 补对角参考线与阴影说明')
d.save(DOCX)

# ── 2. 同名替换媒体字节 ──
with open(PNG, 'rb') as f:
    new_bytes = f.read()
assert Image.open(io.BytesIO(new_bytes)).size == FIG1_SIZE
new_sha = hashlib.sha256(new_bytes).hexdigest()[:12]
tmp = DOCX + '.tmp'
zin = zipfile.ZipFile(DOCX, 'r')
zout = zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED)
for item in zin.infolist():
    data = zin.read(item.filename)
    if item.filename == MEDIA:
        data = new_bytes
    zout.writestr(item, data)
zin.close(); zout.close()
os.replace(tmp, DOCX)
print('%s 已替换，新图 sha256=%s' % (MEDIA, new_sha))

# ── 3. 校验 ──
z = zipfile.ZipFile(DOCX)
assert hashlib.sha256(z.read(MEDIA)).hexdigest()[:12] == new_sha
assert Image.open(io.BytesIO(z.read(MEDIA))).size == FIG1_SIZE
z.close()
d2 = Document(DOCX)
t = None
for p in d2.paragraphs:
    if p.text.strip().startswith('图 1｜'):
        t = p.text; break
assert t is not None
assert '虚线折线为仅具 3–4 年记录的风场' in t
assert '对角虚线为装机份额与新增电量份额均匀对应的参考线' in t
# 8-30 已定的 v6 口径不得被破坏
assert '实心红点为主判据下越线的 4 个风场' in t
assert '两幅区域地图' in t and '三幅区域地图' not in t
assert '5 个越线风场的逐年增益' in t and '6 个越线风场' not in t
print('校验通过：图 1 已换为 (d)(e) 优化版，图注同步且 v6 口径未变')
print('（备份: %s）' % os.path.basename(BAK))
