# -*- coding: utf-8 -*-
"""图 1 重建版换入 v4.9 + 图注同步（2026-08-30）
====================================================================
1. 备份 v4.9 → ADAPEN_manuscript_zh_v4.9_旧图1备份.docx（改动前原件）。
2. 图注第 16 段同步 v6 口径：
   - (b) 实心红点 5→4 个（空心红圈 1 个不变）；
   - (c) 三幅→两幅区域地图（美国东海岸插图已删，F160 掉线）；
   - (d) 6→5 条折线（US East Coast 删除）。
3. 同 rId 替换 word/media/image9.png 字节为新图（4342×5286 与原图同
   像素尺寸，文档声明的显示范围不变、无比例失真）。
4. 输出校验：新图 sha256、图注新旧断言全部打印。
v4.8 原件不动。
"""
import os, sys, io, shutil, hashlib, zipfile
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

HERE = os.path.dirname(os.path.abspath(__file__))
DOCX = os.path.join(HERE, 'ADAPEN_manuscript_zh_v4.9.docx')
BAK = os.path.join(HERE, 'ADAPEN_manuscript_zh_v4.9_旧图1备份.docx')
PNG = os.path.join(HERE, 'figures-nat', 'Fig1_senior_nat.png')

# ── 0. 备份 ──
shutil.copy(DOCX, BAK)
print('备份:', os.path.basename(BAK))

# ── 1. 图注第 16 段 ──
d = Document(DOCX)
P = d.paragraphs
CHANGES = []

def para_replace(p, old, new, tag):
    texts = [r.text for r in p.runs]
    full = ''.join(texts)
    assert old in full, f'{tag}: 未找到 {old[:40]!r}'
    i0 = full.index(old); i1 = i0 + len(old)
    pos = 0; si = None
    for k, t in enumerate(texts):
        if pos <= i0 < pos + len(t):
            si = k; break
        pos += len(t)
    assert si is not None
    prefix = texts[si][:i0 - pos]
    pos_k = pos; last = si
    while pos_k < i1:
        pos_k += len(texts[last]); last += 1
    suffix = texts[last - 1][i1 - (pos_k - len(texts[last - 1])):] if i1 - (pos_k - len(texts[last - 1])) > 0 else ''
    p.runs[si].text = prefix + new + suffix
    for j in range(si + 1, last):
        p.runs[j].text = ''
    CHANGES.append(f'{tag}: {old[:30]}… → {new[:30]}…')

cap = None
for p in P:
    if p.text.strip().startswith('图 1｜'):
        cap = p; break
assert cap is not None, '未找到图 1 图注'
assert '实心红点为主判据下越线的 5 个风场' in cap.text
assert '三幅区域地图' in cap.text
assert '6 个越线风场的逐年增益' in cap.text
para_replace(cap, '实心红点为主判据下越线的 5 个风场',
             '实心红点为主判据下越线的 4 个风场', '图注(b)')
para_replace(cap, '（c）三幅区域地图给出越线风场（星号）',
             '（c）两幅区域地图给出越线风场（星号）', '图注(c)')
para_replace(cap, '（d）折线给出 6 个越线风场的逐年增益',
             '（d）折线给出 5 个越线风场的逐年增益', '图注(d)')
d.save(DOCX)
print('图注改动 %d 处：' % len(CHANGES))
for c in CHANGES:
    print('  -', c)

# ── 2. 同 rId 替换 image9.png 字节 ──
with open(PNG, 'rb') as f:
    new_bytes = f.read()
new_sha = hashlib.sha256(new_bytes).hexdigest()[:12]
tmp = DOCX + '.tmp'
zin = zipfile.ZipFile(DOCX, 'r')
zout = zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED)
for item in zin.infolist():
    data = zin.read(item.filename)
    if item.filename == 'word/media/image9.png':
        data = new_bytes
    zout.writestr(item, data)
zin.close(); zout.close()
os.replace(tmp, DOCX)
print('image9.png 已替换，新图 sha256=%s' % new_sha)

# ── 3. 校验 ──
z = zipfile.ZipFile(DOCX)
assert hashlib.sha256(z.read('word/media/image9.png')).hexdigest()[:12] == new_sha
assert 'word/media/image9.png' in z.namelist()
d2 = Document(DOCX)
t16 = None
for p in d2.paragraphs:
    if p.text.strip().startswith('图 1｜'):
        t16 = p.text; break
assert t16 is not None
assert '实心红点为主判据下越线的 4 个风场' in t16
assert '空心红圈为放宽至 3 年后新增的 1 个' in t16
assert '两幅区域地图' in t16 and '三幅区域地图' not in t16
assert '5 个越线风场的逐年增益' in t16 and '6 个越线风场' not in t16
assert '橙点标出增益最高的 5 个风场' in t16
# 正文关键句复核（v4.9 早前已更新，此处仅确认未被破坏）
for p in d2.paragraphs:
    if '共有五个越线项目' in p.text:
        assert '美国东海岸 F160 的比值降至 0.96' in p.text
print('校验通过：v4.9 图 1 已换为新图，图注与 v6 口径一致')
print('（备份: %s）' % os.path.basename(BAK))
