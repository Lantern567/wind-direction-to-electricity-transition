# -*- coding: utf-8 -*-
"""Assemble the geometry-over-alignment analysis material (Markdown -> Word .docx)
with the 5 figures embedded in-line at their results subsections."""
import os, re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = r"d:/onedrive/01_科研与论文/08_风向建设/分析材料_几何主导杠杆"
MD   = os.path.join(BASE, "分析材料_几何胜于对齐.md")
FIG  = os.path.join(BASE, "figures")
OUT  = os.path.join(BASE, "分析材料_几何胜于对齐.docx")

# subsection number -> (figure file, short in-text caption)
FIGMAP = {
    "2.1": ("Fig1_wakeloss_intensity.png",
            "图 1｜真实海上风场的尾流损失强而系统。(a) 全球 171 场分布，标记大小∝机组数、颜色=平均尾流损失；(b) 场-年尾流损失分布（Gauss, n=1203）；(c) 分区域容量因子与尾流损失；(d) 三尾流模型平均尾流损失，排序稳定。"),
    "2.2": ("Fig2_geometry_drives_wake.png",
            "图 2｜尾流的场内主导内因是几何拥挤而非朝向失配。(a) 打乱对照：真实排布 vs 同机组随机散布的尾流损失；(b) 真实最近邻间距分布（n=171，单位 D）与 5.1D/7D 参考线；(c) 间距↔尾流损失散点及拟合线（Spearman ρ=−0.25, p=0.001）。"),
    "2.3": ("Fig3_orientation_controllability.png",
            "图 3｜刚性旋转到历史最优朝向的收益幅度小但高度稳健。(a) AEP 提升分布（n=1203，均值+0.9%、胜率64.3%、p=3e-44）；(b) 实际朝向与最优朝向偏差绝对值分布（n=167，中位51.2°）；(c) 分国家平均 AEP 提升；(d) 代表场整场刚性旋转的尾流效率响应。"),
    "2.4": ("Fig4_recoverable_envelope.png",
            "图 4｜可回收价值几乎全在点阵而非朝向。(a) 按杠杆的可回收 AEP：朝向刚性旋转 +0.9%（实测）、点阵优化 POC（本研究，中位 +6.2%、均值 +12.2%）、理想 5D 方阵优化上限（W&P 2024, 6–7%）；(b) 真实最近邻间距分布 vs 理想 5D 基线（84.8% 更密）；(c) 12 个代表场点阵优化可回收 AEP vs 真实间距，标记大小∝机组数，越密回收越大（最密两场 >40%）。"),
    "2.5": ("Fig5_densification_spillover.png",
            "图 5｜密集化把场内尾流推向跨场外部性。(a) 不同邻近阈值下的相邻风场对数；52/171 场（30%）有 <10km 近邻；(b) 邻近场簇跨场仿真：东亚对（12.9km）独立 vs 合并域，跨场尾流 +0.47%（保守下界，欧洲对因单点风资源混淆排除）；(c) 技术×密集化情景：三档机型（IEA 10/15/22MW）平均尾流损失与 CF，10→22MW 尾流 +4.3pp、CF −1.8pp。"),
    "2.7": ("Fig7_geographic_three_levers.png",
            "图 7｜三因素的逐站点地理图景（同一把 ±20pp 发散标尺：红=损失、蓝=增益、白=零；点大小∝机组数；上行欧洲、下行东亚）。(a) 风速禀赋（vs 全样本平均风况）：北海 +5~+15pp、东亚季风区 −5~−15pp、台湾海峡 +19pp；(b) 朝向罚分（vs 自身最优）：中位 −0.3pp 几乎全图空白，唯越南集群 −18pp；(c) 点阵/尾流罚分（vs 无尾流理想）：171 场全为负，中位 −11pp、最深 −48pp。美东 5 场因幅面未绘。"),
    "2.6": ("Fig6_windspeed_role.png",
            "图 6｜风速的角色：发电的天花板与噪声，而非尾流的结构性解释。(a) 按风速分箱的尾流损失：171 场观测（中位/IQR）+ 受控扫描（8×8 5D 方阵、真实密排场）；灰柱=能量占比；虚线=额定 ~11m/s；(b) 归因防守（n=72）：控间距后 WPD↔尾流 −0.25→−0.00 塌缩，控风资源后间距 −0.52→−0.47 存活；(c) CF vs 年均轮毂风速（n=171，r=0.92，+0.074 CF/(m/s)）；(d) 风年噪声 5.2% vs 朝向 0.9% vs 几何 6.2%——仅几何越过噪声线。"),
}

doc = Document()
# base style
st = doc.styles["Normal"]
st.font.name = "Times New Roman"
st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(__import__("docx").oxml.ns.qn("w:eastAsia"), "宋体")

BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
ITAL_RE = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")


def add_runs(par, text):
    """Render **bold** and *italic* inline markup into a paragraph."""
    # split on bold first
    idx = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > idx:
            par.add_run(text[idx:m.start()])
        r = par.add_run(m.group(1)); r.bold = True
        idx = m.end()
    if idx < len(text):
        par.add_run(text[idx:])


def insert_figure(subsec):
    if subsec not in FIGMAP:
        return
    fname, cap = FIGMAP[subsec]
    fpath = os.path.join(FIG, fname)
    if not os.path.exists(fpath):
        return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(fpath, width=Inches(6.3))
    cp = doc.add_paragraph()
    cr = cp.add_run(cap); cr.font.size = Pt(8.5); cr.italic = True
    cp.paragraph_format.space_after = Pt(8)


lines = open(MD, encoding="utf-8").read().split("\n")
i = 0
cur_subsec = None
while i < len(lines):
    ln = lines[i].rstrip()
    if ln.startswith("# "):
        h = doc.add_heading("", level=0); add_runs(h, ln[2:])
    elif ln.startswith("### "):
        h = doc.add_heading("", level=2); add_runs(h, ln[4:])
        m = re.match(r"###\s+(\d\.\d)", ln)
        # results subsection -> remember to insert its figure after its body
        cur_subsec = None
        m2 = re.match(r"###\s+(2\.\d)\s", ln)
        if m2:
            cur_subsec = m2.group(1)
    elif ln.startswith("## "):
        h = doc.add_heading("", level=1); add_runs(h, ln[3:])
    elif ln.startswith("#### "):
        h = doc.add_heading("", level=3); add_runs(h, ln[5:])
    elif ln.startswith("> "):
        p = doc.add_paragraph(style="Intense Quote"); add_runs(p, ln[2:])
        # a "结论 N｜" quote closes a results subsection -> drop the figure right after
        if cur_subsec and ("结论" in ln):
            insert_figure(cur_subsec); cur_subsec = None
    elif ln.startswith("!["):
        pass  # markdown image line -- real image inserted via FIGMAP
    elif ln.startswith("*图") and ln.rstrip().endswith("*"):
        pass  # markdown caption line -- caption inserted via FIGMAP
    elif ln.startswith("- "):
        p = doc.add_paragraph(style="List Bullet"); add_runs(p, ln[2:])
    elif re.match(r"^\d+\.\s", ln):
        p = doc.add_paragraph(style="List Number"); add_runs(p, ln[ln.index(".")+2:])
    elif ln.startswith("|") or ln.startswith("---") or ln.startswith("> 版本") or ln.startswith("> 说明"):
        # skip md table separators / metadata banners
        if ln.startswith("> "):
            p = doc.add_paragraph(); r = p.add_run(ln[2:]); r.italic = True; r.font.size = Pt(8.5)
    elif ln.strip() == "":
        pass
    elif ln.startswith("*") and ln.endswith("*") and "Geometry" in ln:
        p = doc.add_paragraph(); r = p.add_run(ln.strip("*")); r.italic = True
    else:
        add_runs(doc.add_paragraph(), ln)
    i += 1

doc.save(OUT)
print("saved:", OUT, "bytes:", os.path.getsize(OUT))
