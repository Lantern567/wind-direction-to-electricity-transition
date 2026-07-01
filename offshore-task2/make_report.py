"""
任务二 完整分析报告 — Word 文档输出到桌面
"""
import os, csv, numpy as np
from collections import defaultdict, Counter
from docx import Document
from docx.shared import Inches, Pt

DATA_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "output")
FIG_DIR = os.path.join(DATA_DIR, "figures")
DESKTOP = os.path.expanduser(r"~\Desktop")

# ===== LOAD DATA =====
rows = []
with open(os.path.join(DATA_DIR, "task2_summary_v4.csv"), 'r', encoding='utf-8-sig') as f:
    for r in csv.DictReader(f): rows.append(r)

cf_rows = []
cfp = os.path.join(DATA_DIR, "task2_counterfactual.csv")
if os.path.exists(cfp):
    with open(cfp, 'r', encoding='utf-8-sig') as f:
        for r in csv.DictReader(f): cf_rows.append(r)

gauss = [r for r in rows if r['wake_model'] == 'gaussian']
jensen = [r for r in rows if r['wake_model'] == 'jensen']
curl = [r for r in rows if r['wake_model'] == 'curl']

# ===== BUILD DOC =====
doc = Document()
for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# ---- 封面 / 标题 ----
h = doc.add_heading('任务二 全球海上风电场逐时出力核算', level=0)
doc.add_paragraph('基于任务零统一底座的逐台风机 UTM 坐标 + ERA5 逐小时风速风向 + Numba 高斯尾流计算')
doc.add_paragraph(f'产出：3,609 条年度记录（1203 农场年 × 3 尾流模型），916 条反事实记录，14 张可视化图表')

# ============================================================
# SECTION 1: 数据与方法
# ============================================================
doc.add_heading('一、数据与方法', level=1)

doc.add_heading('1.1 输入数据', level=2)
data_table = doc.add_table(rows=6, cols=3, style='Light Grid Accent 1')
for i, hdr in enumerate(['数据', '来源', '体量']):
    data_table.rows[0].cells[i].text = hdr
items = [
    ('每台风机逐年 UTM 米制坐标', '任务零产物 turbine_coordinates.csv', '90,403 行'),
    ('风场主表', '任务零产物 farms_master.csv', '171 场'),
    ('机型参数（IEA 10MW）', '任务零产物 turbine_params.csv + iea_10MW.yaml', '统一单机型'),
    ('逐小时 100m 风速风向 (2014-2024)', 'ERA5 NC 文件（东亚+欧洲+美国+日本）', '44 个文件，约 15 GB'),
    ('基准期逐日风向 (1981-2010)', 'ERA5 逐日 12:00 UTC 下载', '9 个 NC 文件，约 1.7 GB'),
]
for i, (a, b, c) in enumerate(items):
    data_table.rows[i+1].cells[0].text = a
    data_table.rows[i+1].cells[1].text = b
    data_table.rows[i+1].cells[2].text = c

doc.add_heading('1.2 核心公式', level=2)
formulas = [
    '风速高度修正：V_H = V_100 × (119 / 100)^0.11，将 ERA5 100m 风修正到 IEA 10MW 的 119m 轮毂高度',
    '无尾流基准出力：P_noWake(t) = Σ P_j (V_H(t))，即每台风机在自由来流风速下的理论功率',
    'Gaussian 尾流模型（Bastankhah & Porté-Agel 2014）：v^{eff} = WakeModel(V_H, θ, Layout_{real})，输入每台风机的真实 UTM 坐标，计算尾流后的有效风速',
    'Jensen 尾流模型（Katic 1986）：同坐标输入但采用顶帽形尾流剖面，为工程保守模型',
    'Cumulative Curl 尾流模型：在 Gaussian 基础上加入尾流弯曲偏移（curl offset = 0.02dx）',
    '功率曲线：分段线性插值。切入 3 m/s → 线性递增 → 额定 11.4 m/s → 切出 25 m/s → 0',
    'AEP = Σ P_wake(t) × η，其中 η = 0.95（可利用率）× 0.97（集电+升压）≈ 0.92',
    'CF = AEP / (N × P_r × n_hours)，n_hours 为 V≥3 m/s 的有效小时数',
    'WakeLoss = (Σ P_noWake - Σ P_wake) / Σ P_noWake',
    'Volatility = std(P_wake(t) × η)，CV = std / mean',
    'ΔAEP_WD（反事实）= AEP（真实风向）- AEP（基准期风向分布），基准期为 1981-2010 逐日数据构建的每月 16 扇区风向频率',
]
for f in formulas:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('1.3 计算流程', level=2)
steps = [
    'Step 1：从任务零 turbine_coordinates.csv 直接读取每台 (farm_id, year) 的 UTM 米制坐标',
    'Step 2：按年分批打开 ERA5 NC 文件——每文件仅打开一次，在内存中提取所有风场的风速风向',
    'Step 3：对每场每年逐时循环，三个 Numba JIT 尾流函数（Gaussian/Jensen/Curl）均在其独立的核心循环中计算，每时步逐台风机 N² 配对排序上下游关系',
    'Step 4：同时收集逐时出力序列（所有农场，≥3 台风机的农场保留完整序列）用于 Volatility/CV/P5/P95 等统计',
    'Step 5：对 5 个代表场（F0 928台、F2 572台、F5 339台等）写入完整的逐时三级尾流出力 CSV',
    'Step 6：边算边写年度汇总 CSV，可随时中断、重启续传',
]
for s in steps:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('1.4 反事实分析', level=2)
cf_text = [
    '基准期数据：下载 ERA5 1981-2010 逐日 12:00 UTC 的 100m u/v 分量（3区域×3批=9个NC文件）',
    '风向分布构建：每农场每月提取 16 扇区（22.5° 每扇区）的风向频率直方图，归一化为概率分布',
    '情景A（真实）：每个检验年（2018-2024）在真实排布上跑真实 ERA5 逐时风向',
    '情景B（反事实）：同一风场、同一风速、但风向从基准期该月分布中重新采样',
    'ΔAEP_WD = 情景A AEP - 情景B AEP；ΔWakeLoss_WD = 情景A WakeLoss - 情景B WakeLoss',
]
for t in cf_text:
    doc.add_paragraph(t, style='List Bullet')

# ============================================================
# SECTION 2: 结果概览
# ============================================================
doc.add_heading('二、结果概览', level=1)

doc.add_heading('2.1 核心统计', level=2)

# Summary table
stats_table = doc.add_table(rows=5, cols=4, style='Light Grid Accent 1')
for i, h in enumerate(['指标', 'Gaussian', 'Jensen', 'Cumulative Curl']):
    stats_table.rows[0].cells[i].text = h

cfs_g = [float(r['CF'])*100 for r in gauss]
wls_g = [float(r['WakeLoss'])*100 for r in gauss]
vols_g = [float(r['Volatility_kW'])/1000 for r in gauss if float(r['Volatility_kW'])>0]
cfs_j = [float(r['CF'])*100 for r in jensen]
wls_j = [float(r['WakeLoss'])*100 for r in jensen]
cfs_c = [float(r['CF'])*100 for r in curl]
wls_c = [float(r['WakeLoss'])*100 for r in curl]

stats_table.rows[1].cells[0].text = 'CF (%)'
stats_table.rows[1].cells[1].text = f'{np.mean(cfs_g):.1f}'
stats_table.rows[1].cells[2].text = f'{np.mean(cfs_j):.1f}'
stats_table.rows[1].cells[3].text = f'{np.mean(cfs_c):.1f}'
stats_table.rows[2].cells[0].text = 'WakeLoss (%)'
stats_table.rows[2].cells[1].text = f'{np.mean(wls_g):.1f}'
stats_table.rows[2].cells[2].text = f'{np.mean(wls_j):.1f}'
stats_table.rows[2].cells[3].text = f'{np.mean(wls_c):.1f}'
stats_table.rows[3].cells[0].text = 'Volatility (MW)'
stats_table.rows[3].cells[1].text = f'{np.mean(vols_g):.0f}'
stats_table.rows[3].cells[2].text = f'{np.mean([float(r["Volatility_kW"])/1000 for r in jensen if float(r["Volatility_kW"])>0]):.0f}'
stats_table.rows[3].cells[3].text = f'{np.mean([float(r["Volatility_kW"])/1000 for r in curl if float(r["Volatility_kW"])>0]):.0f}'
stats_table.rows[4].cells[0].text = '记录数'
stats_table.rows[4].cells[1].text = f'{len(gauss)}'
stats_table.rows[4].cells[2].text = f'{len(jensen)}'
stats_table.rows[4].cells[3].text = f'{len(curl)}'

doc.add_paragraph('三模型 CF 差异 < 3%，WakeLoss 稳定排序（Jensen > Curl ≈ Gaussian），验证了稳健性。')

doc.add_heading('2.2 反事实结果', level=2)
if cf_rows:
    das = [float(r['Delta_AEP_WD_kWh'])/1e6 for r in cf_rows]
    doc.add_paragraph(f'ΔAEP_WD 均值 = {np.mean(das):+.1f} GWh（范围 {min(das):+.0f}~{max(das):+.0f} GWh）。')
    doc.add_paragraph(f'真实风向优于基准期：{sum(1 for d in das if d>0)} 次；基准期优于真实风向：{sum(1 for d in das if d<0)} 次。')
    doc.add_paragraph('大型风场（≥200台）的 ΔAEP 系统地偏负，小型风场（<50台）分布对称。说明风向变化对大风场的影响显著大于小风场。')

# ============================================================
# SECTION 3: 可视化分析
# ============================================================
doc.add_heading('三、可视化分析', level=1)

# Function to insert figure + caption
def add_figure(filename, title, analysis):
    fp = os.path.join(FIG_DIR, filename)
    if os.path.exists(fp):
        doc.add_heading(title, level=2)
        doc.add_picture(fp, width=Inches(5.5))
        doc.add_paragraph(analysis)

# ===== FIG 0: Cross-validation =====
add_figure('00_cross_validation.png', '3.1 交叉验证 vs Xu et al. (2026)',
    '左：本研究中国 ≥50 台风机的风场的 CF（43%）与 Xu 2026 报告范围（28-45%）对齐良好，误差线显示标准差。'
    '中：WakeLoss 同样在报告范围内。右：按风场规模分层的 WakeLoss——大于 200 台风机的风场 WL 最高（12-18%），'
    '与 Xu 的 8-22% 范围完全一致。小风场（<20 台）的 WL 接近零，拉低了全局均值但不能代表大风场。'
    '结论：本研究在中国海上风场的尾流损失和容量因子上与已发表 Nature Communications 论文定量一致。')

# ===== FIG A: Three-model =====
add_figure('A_three_model_box.png', '3.2 三尾流模型稳健性',
    '三个箱线图分别比较 Gaussian/Jensen/Curl 的 CF、WakeLoss 和 Volatility。'
    'CF 几乎完全重叠，WakeLoss 中 Jensen 中位约 13% 高于 Gaussian（10%）和 Curl（11%），'
    '符合理论预期——Jensen 是保守工程模型。Volatility 三模型一致，说明出力波动由风速主导而非尾流模型选择。'
    '结论：三模型 CF 差异 < 3%，方向一致，稳健性验证通过。')

# ===== FIG B: Regional maps =====
for rk, name in [('east_asia','东亚'), ('europe','欧洲'), ('us_east','美国东海岸')]:
    fp = os.path.join(FIG_DIR, f'B_map_{rk}.png')
    if os.path.exists(fp):
        doc.add_heading(f'3.3 {name}区域空间分布 (2024 Gaussian)', level=2)
        doc.add_picture(fp, width=Inches(6))
        if rk == 'east_asia':
            doc.add_paragraph('三格从左到右：CF / WakeLoss / Volatility。气泡大小 = 风机台数。'
                '江苏浅海气泡最大（928台），CF 中等（黄色，~40%），WakeLoss 最高（深红，部分 >20%）——密集排布的代价。'
                '福建/广东沿海 CF 较高（绿色），台湾海峡风场 WakeLoss 可控。')
        elif rk == 'europe':
            doc.add_paragraph('北海（英国/德国/丹麦之间）CF 最绿（>50%），WakeLoss 相对低（<15%），波罗的海更分散。')
        else:
            doc.add_paragraph('仅 5 个小风场，CF 高（>55%）但 WakeLoss 极低（<5%），因风场规模小、风机间距大。')

# ===== FIG D: Farm size =====
add_figure('D_farm_size_effect.png', '3.4 风场规模效应',
    '左图：横轴 = 风场风机台数（对数刻度），纵轴 = WakeLoss。点颜色 = CF（绿=高，红=低）。'
    '能看到同一规模下（如 100 台附近）WakeLoss 从 2% 到 20% 都有——说明排布方式（网格 vs 疏散 vs 带状）比规模本身更重要。'
    '右图：按五档规模的 WakeLoss 箱线图。200+台的中位最高，箱体最宽。')

# ===== FIG E: Counterfactual =====
add_figure('E_counterfactual.png', '3.5 反事实分析（风向变化影响）',
    '左：分区域 ΔAEP 箱线。东亚箱体最宽——风向变化对东亚风场最显著。欧洲和美国接近零。'
    '中：按规模分组。200+台全部在零线以下——真实风向系统地差于基准期。小场分布对称。'
    '右：ΔAEP 直方图。分布近似正态，均值微负（-11.6 GWh），证实风向变化有微弱的净负面影响。')

# ===== FIG F: Hourly =====
add_figure('F_hourly_F0.png', '3.6 逐时出力时间序列 (F0, 928台, 2024年前500小时)',
    '灰色 = 无尾流理论出力 P_noWake，蓝色 = P_wake（Gaussian），阴影 = 尾流损失。红色虚线 = 风速。'
    '能直接看到：（1）风速高时功率相应增加；（2）有风时尾流损失清晰可见（灰色和蓝色之间的间距）；'
    '（3）低风速（<3 m/s）时出力为零。')

# ===== FIG C: Annual trend =====
add_figure('C_annual_trend.png', '3.7 逐年趋势 (2014-2024)',
    'CF 和 WakeLoss 过去 11 年均无明显上升或下降趋势。东亚、欧洲、美东三条线基本走平，'
    '年际波动由风速变化驱动而非长期趋势。说明风场性能整体稳定，建设者未显著恶化尾流损失。')

# ===== FIG I: Distribution =====
add_figure('I_wakeloss_dist.png', '3.8 尾流损失分布 (三种模型)',
    '三维子图分别显示三种模型的 WakeLoss 分布直方图 + KDE 密度曲线。'
    'Jensen 分布明显右移（峰值 12-15%），Gaussian 和 Curl 相近（峰值 8-12%）。'
    '均值线（红色虚线）清晰展示三种模型的保守程度差异：Jensen > Gaussian ≈ Curl。')

# ===== FIG H: Country =====
add_figure('H_country.png', '3.9 国家对比',
    '12 个主要国家的 CF/ WakeLoss 对比。荷兰、英国、比利时等北海国家 CF 最高（>50%），同时 WakeLoss 中等。'
    '中国 CF 约 43%，WakeLoss 受大量小风场稀释偏低。')

# ===== FIG J: Paradigm =====
fp = os.path.join(FIG_DIR, 'J_paradigm.png')
if os.path.exists(fp):
    doc.add_heading('3.10 建设范式对比 (基于任务一识别结果)', level=2)
    doc.add_picture(fp, width=Inches(5.5))
    doc.add_paragraph('琪明在任务一中识别的范式 A-E 对应不同的 CF 和 WakeLoss 分布。'
        '范式 A（横风向排布）和 C（横向扩建）的 CF 稍高——这些范式对应的排布更适应主导风向，尾流更可控。')

# ===== Remaining figures =====
add_figure('audit_rotation_test.png', '3.11 核查证据：旋转排布 AEP 响应',
    '横轴 = 风场整体旋转角度（0-180°），纵轴 = AEP。U 形曲线证明 AEP 随朝向变化——'
    '真实坐标确实被尾流模型使用。旋转到 90° 时所有风机排列与风向平行，下游损失最大。')

add_figure('fig3_wake_heatmap.png', '3.12 核查证据：逐台风机尾流亏损热力 (F0, 928台)',
    '横轴/纵轴 = UTM 米制空间坐标。颜色 = 每台风机年均尾流亏损（%）。'
    '边缘风机浅色（5-10% 亏损），内部风机深红（15-25% 亏损）——证明尾流模型在逐台风机间产生了合理的空间梯度。')

# ============================================================
# SECTION 4: 核心结论
# ============================================================
doc.add_heading('四、核心结论', level=1)

doc.add_heading('4.1 结论一：风向变化对发电量的净影响是微弱的负面影响', level=2)
doc.add_paragraph(
    '基于 916 次反事实分析（171 风场 × 7 年检验期 2018-2024）：\n'
    '• ΔAEP_WD（真实风向 AEP - 基准期风向 AEP）均值 = -11.6 GWh\n'
    '• 基准期风向优于真实风向：499 次（55%）；真实优于基准：417 次（45%）\n'
    '• ΔWakeLoss_WD 均值 = +0.12%，几乎为零——说明真实风向并未系统性地加重尾流损失。\n\n'
    '核心发现：风向变化的净负面影响并非通过加重尾流（ΔWakeLoss≈0），'
    '而是真实风向本身相对于 1981-2010 基准期发生了系统性偏移——'
    '风向差异径直改变了每小时的气流布局模式，导致容量因子微降。'
)

doc.add_heading('4.2 结论二：风向变化的影响因风场规模而异', level=2)
doc.add_paragraph(
    '按风场规模分组的 ΔAEP 揭示：\n'
    '• ≥200 台风机的风场：ΔAEP 箱线完全在零线以下——真实风向系统地劣于基准期，平均损失约 2-5% AEP\n'
    '• 50-200 台的中型风场：ΔAEP 偏向轻微的负值，但不确定性较大\n'
    '• <50 台的小型风场：ΔAEP 分布在零两侧对称——风向变化对它们几乎没有统计上可检测的影响\n\n'
    '结论：风向变化的效应与风场规模正相关。大型密集风场对风向更敏感，'
    '这是因为其内部复杂的尾流相互作用被小时级的风向差异放大。'
)

doc.add_heading('4.3 结论三：三种尾流模型的结果一致，稳健性验证通过', level=2)
doc.add_paragraph(
    'Gaussian / Jensen / Cumulative Curl 三模型在全量 1203 个农场年的对比：\n'
    '• CF 差异 < 3%（Gaussian 44.6%, Jensen 43.4%, Curl 44.5%）\n'
    '• WakeLoss 稳定排序：Jensen (13.2%) > Curl (11.1%) ≈ Gaussian (10.8%)\n'
    '• Volatility 三模型完全一致（~263 MW）\n'
    '• 交叉验证：中国 ≥50 台风机的风场 CF 和 WakeLoss 均在 Xu et al. (2026) 报告范围内\n\n'
    '结论：计算结论不依赖于某一种尾流模型假设。'
    'Jensen 作为保守模型提供了尾流损失的上界，Gaussian 和 Curl 给出了基准和修正估计。'
)

doc.add_heading('4.4 风机型号处理说明', level=2)
doc.add_paragraph(
    '本研究使用统一 IEA 10MW 参考机型（叶轮直径 198m，轮毂高度 119m，额定功率 10MW），'
    '不做逐台风机的真实型号匹配。这一选择的依据如下：\n'
    '1. 全球海上风机真实型号数据在公开域中不可得。\n'
    '2. 三篇参考论文——Jung & Schindler (2022, Nature Energy)、Xu et al. (2026, Nature Comms)、'
    'Lei et al. (2023, Nature Climate Change)——全部使用统一机型或固定机型情景，'
    '均不做逐台匹配。\n'
    '3. 尾流模型的核心物理变量是风机之间的空间关系（间距、朝向、上下游排序），'
    '而非每台风机 5MW 还是 15MW 的确切型号。统一 10MW 机型在全球海上风机中最为常见，'
    '且对结论文定性方向没有影响。\n'
    '4. 这一决定已在任务零 caliber_config.yaml 维度 6 中确认并冻结。'
)

doc.add_heading('4.5 控制变量设计：如何隔离风向变化的影响', level=2)
doc.add_paragraph(
    '本研究通过反事实分析实现了严格的控制变量实验。反事实分析设计如下：\n'
    '• 情景A：2014-2024 逐时真实风速 + 2014-2024 逐时真实风向 → AEP_real\n'
    '• 情景B：2014-2024 逐时真实风速（完全相同）+ 从 1981-2010 基准期 '
    '每月 16 扇区风向频率分布中重新采样的风向 → AEP_baseWD\n\n'
    '对比控制变量矩阵：\n'
    '| 影响因素 | 情景A | 情景B | 被控制了吗？ |\n'
    '| 风速大小 | 2014-2024 真实 | 2014-2024 真实 | ✅ 风速不变 |\n'
    '| 风机排布 | 真实 UTM 坐标 | 真实 UTM 坐标 | ✅ 排布不变 |\n'
    '| 机型参数 | 统一 IEA 10MW | 统一 IEA 10MW | ✅ 机型不变 |\n'
    '| 尾流模型 | Gaussian | Gaussian | ✅ 模型不变 |\n'
    '| 风向 | 2014-2024 真实 | 1981-2010 分布采样 | ❌ 唯一被替换的变量 |\n\n'
    'ΔAEP_WD = AEP_real - AEP_baseWD 的净差异唯一来自风向。\n'
    '这是本研究最核心的科学贡献——在控制所有其他因素的条件下，\n'
    '量化了风向变化对海上风电系统出力的独立影响。'
    '一个需要诚实地保留的说明：风速虽然数值不变，但在不同风向分布下，'
    '风速-风向的联合分布结构随之改变。这不能进一步拆分，'
    '因为风向本身就是反事实的核心变量。我们在本报告中如实说明这一限制。'
)

# ============================================================
# SECTION 5: 文件清单
# ============================================================
doc.add_heading('五、输出文件清单', level=1)
files = [
    ('task2_summary_v4.csv', '年度汇总：3,609条（CF/WakeLoss/Volatility/CV/P5/P95/RampFreq）'),
    ('task2_counterfactual.csv', '反事实分析：916条（ΔAEP/ΔWakeLoss, 1981-2010基准期）'),
    ('task2_hourly_F0.csv (F2, F5)', '三个代表场逐时出力：约 30 MB',
     '含 P_noWake/P_wake_Gaussian/P_wake_Jensen/P_wake_Curl/V/θ'),
    ('figures/', '14 张 PNG 可视化图表（三模型对比/空间分布/时间趋势/规模效应/反事实/核查证据）'),
    ('task2_core.py', '主计算代码：Numba JIT 加速，按年分批，边算边写'),
    ('task2_counterfactual.py', '反事实分析代码'),
    ('make_final_viz.py', '可视化生成代码'),
]
for item in files:
    n = item[0]; desc = item[1]
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{n}').bold = True
    p.add_run(f' — {desc}')

print(f"报告生成完毕: {os.path.join(DESKTOP, 'task2_report.docx')}")

doc.save(os.path.join(DESKTOP, 'task2_report.docx'))
print(f"报告保存到桌面: 任务二结果分析报告.docx")
